"""报告生成：LLM 分析 + 可视化图表 + 专业 DOCX 导出"""
import os
import re
import io
import threading
import numpy as np
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties, fontManager, findfont

from app.config.settings import get_settings
from app.config.logging import get_logger

logger = get_logger("services.report")
settings = get_settings()

# Matplotlib 全局锁（模块级 rcParams/fontManager 非线程安全）
_mpl_lock = threading.Lock()

# 中文字体 — 跨平台候选列表
_FONT_CANDIDATES = [
    # Windows
    ("C:/Windows/Fonts/msyh.ttc", "Microsoft YaHei"),
    ("C:/Windows/Fonts/simsun.ttc", "SimSun"),
    ("C:/Windows/Fonts/simhei.ttf", "SimHei"),
    # Linux
    ("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc", "WenQuanYi Zen Hei"),
    ("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc", "WenQuanYi Micro Hei"),
    ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", "Noto Sans CJK SC"),
    ("/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf", "Droid Sans Fallback"),
    # macOS
    ("/System/Library/Fonts/PingFang.ttc", "PingFang SC"),
    ("/Library/Fonts/Arial Unicode.ttf", "Arial Unicode MS"),
]

FONT_PATH = None
FONT_NAME = None
for _path, _name in _FONT_CANDIDATES:
    if os.path.exists(_path):
        fontManager.addfont(_path)
        FONT_PATH = _path
        FONT_NAME = _name
        break

# Dynamic fallback: scan system for any CJK font
if FONT_PATH is None:
    try:
        _cjk_fonts = [f for f in fontManager.ttflist if any(
            k in f.name.lower() for k in ('cjk', 'hei', 'song', 'ming', 'noto sans', 'wqy')
        )]
        if _cjk_fonts:
            FONT_PATH = _cjk_fonts[0].fname
            FONT_NAME = _cjk_fonts[0].name
    except Exception:
        pass

if FONT_NAME:
    plt.rcParams['font.sans-serif'] = [FONT_NAME, 'DejaVu Sans']
    plt.rcParams['font.family'] = 'sans-serif'
else:
    plt.rcParams['font.family'] = 'sans-serif'
plt.rcParams['axes.unicode_minus'] = False

INDICATOR_CN = {
    "chlorophyll": "叶绿素 (µg/L)", "dissolved_oxygen": "溶解氧 (mg/L)",
    "temperature": "水温 (°C)", "ph": "pH", "turbidity": "浊度 (NTU)",
}
SHORT_CN = {"chl": "叶绿素", "odo": "溶解氧", "temp": "水温", "ph": "pH", "turb": "浊度"}
COLORS = ["#409eff", "#67c23a", "#e6a23c", "#f56c6c", "#909399"]

PROMPT = """你是一名资深水质分析专家。请根据以下数据撰写专业水质巡航分析报告。

巡航基础信息
- 水库：{reservoir_name}
- 时间：{date}
- 采样点：{point_count} 个，深度层：{depth_layers}

指标统计
- 叶绿素：均值 {chl_mean:.2f} µg/L，异常率 {chl_anomaly_rate:.1%}
- 溶解氧：均值 {odo_mean:.2f} mg/L，异常率 {odo_anomaly_rate:.1%}
- 水温：均值 {temp_mean:.2f} °C，异常率 {temp_anomaly_rate:.1%}
- pH：均值 {ph_mean:.2f}，异常率 {ph_anomaly_rate:.1%}
- 浊度：均值 {turb_mean:.2f} NTU，异常率 {turb_anomaly_rate:.1%}

异常点位（共 {anomaly_count} 个）：
{anomaly_sample}

历史相似案例：{cases_summary}

请按以下 4 个部分撰写（每部分以"##SECTION:标题##"开头，后续为纯文本段落，不要使用任何 Markdown 标记）：

##SECTION:总体水质评价##
（1-2 段：综合各指标评价水库水质状况，指出达标和超标指标，判断水体健康等级）

##SECTION:异常区域分析##
（2-3 段：分析异常点的空间分布特征，结合深度规律，推断可能的污染原因）

##SECTION:历史案例对比##
（1-2 段：与相似历史案例对比，分析水质变化趋势）

##SECTION:改善建议##
（3-5 条具体、可操作的建议，编号列出）

要求：专业、客观、数据驱动，每部分 200-400 字，使用标准中文书面语。"""


def compute_stats(rows, anomalies):
    indicators = ["chlorophyll", "dissolved_oxygen", "temperature", "ph", "turbidity"]
    shorts = ["chl", "odo", "temp", "ph", "turb"]
    stats = {}
    for ind, sn in zip(indicators, shorts):
        vals = [getattr(r, ind) for r in rows if getattr(r, ind) is not None]
        anom = [a for a in anomalies if a.get("indicator") == sn]
        stats[f"{sn}_mean"] = float(np.mean(vals)) if vals else 0
        stats[f"{sn}_std"] = float(np.std(vals)) if vals else 0
        stats[f"{sn}_min"] = float(np.min(vals)) if vals else 0
        stats[f"{sn}_max"] = float(np.max(vals)) if vals else 0
        stats[f"{sn}_anomaly_rate"] = len(anom) / len(rows) if rows else 0
        stats[f"{sn}_count"] = len(vals)
        stats[f"{sn}_anomaly_count"] = len(anom)
    return stats


def build_prompt(task_info, stats, anomalies, similar):
    anomaly_text = "\n".join(
        f"  坐标({a['lon']:.5f}, {a['lat']:.5f}) 深度{a['depth']}m "
        f"指标{SHORT_CN.get(a['indicator'], a['indicator'])}={a['value']:.2f} 方法:{a['method']}"
        for a in anomalies[:15]
    ) or "无异常点"

    if similar:
        cases = "已匹配 " + str(len(similar)) + " 个相似案例"
    else:
        cases = "暂无历史相似案例可供对比"

    return PROMPT.format(
        reservoir_name=task_info.get("reservoir_name", "未知"),
        date=task_info.get("created_at", datetime.utcnow().strftime("%Y-%m-%d")),
        point_count=task_info.get("total_points", 0),
        depth_layers=task_info.get("depth_layers", "N/A"),
        anomaly_count=len(anomalies),
        anomaly_sample=anomaly_text, cases_summary=cases, **stats,
    )


def call_llm(prompt: str) -> str:
    import time
    from openai import OpenAI, APIStatusError, APITimeoutError, APIConnectionError
    import httpx

    if settings.LLM_PROVIDER == "deepseek":
        api_key = settings.DEEPSEEK_API_KEY
        base_url = settings.DEEPSEEK_BASE_URL
        model = settings.DEEPSEEK_MODEL
    else:
        api_key = settings.OPENAI_API_KEY
        base_url = settings.OPENAI_BASE_URL
        model = settings.OPENAI_MODEL

    timeout = httpx.Timeout(60.0, read=300.0, write=30.0, pool=10.0)
    last_error = None

    for attempt in range(3):
        try:
            client = OpenAI(api_key=api_key, base_url=base_url, timeout=timeout)
            resp = client.chat.completions.create(
                model=model, messages=[{"role": "user", "content": prompt}],
                temperature=0.7, max_tokens=2048,
            )
            content = resp.choices[0].message.content
            if not content or not content.strip():
                raise RuntimeError(f"LLM 返回空内容 (model={model})")
            required_sections = ["总体水质评价", "异常区域分析", "历史案例对比", "改善建议"]
            missing = [s for s in required_sections if f"##SECTION:{s}##" not in content]
            if missing:
                logger.warning(f"LLM 响应缺少章节: {missing} | 响应长度: {len(content)}")
            logger.info(f"LLM 响应成功 | 长度: {len(content)} 字符 | 模型: {model} | 尝试: {attempt + 1}")
            return content
        except (APIStatusError, APITimeoutError, APIConnectionError) as e:
            last_error = e
            if attempt < 2:
                wait = 2 ** attempt
                logger.warning(f"LLM 调用失败 (尝试 {attempt + 1}/3)，{wait}s 后重试: {e}")
                time.sleep(wait)
        except RuntimeError:
            raise

    raise last_error or RuntimeError(f"LLM 调用失败 (model={model})")


# ─── 图表生成 ─────────────────────────────────────────────

def _fig_to_bytes(fig):
    buf = io.BytesIO()
    try:
        with _mpl_lock:
            fig.savefig(buf, format='png', dpi=200, bbox_inches='tight')
    except ValueError:
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=200)
    buf.seek(0)
    plt.close(fig)
    return buf


def make_depth_profile_chart(rows, anomalies, task_name):
    """5 指标深度剖面图"""
    indicators = ["chlorophyll", "dissolved_oxygen", "temperature", "ph", "turbidity"]
    shorts = ["chl", "odo", "temp", "ph", "turb"]

    fig, axes = plt.subplots(1, 5, figsize=(14, 5))
    for idx, (ind, sn) in enumerate(zip(indicators, shorts)):
        ax = axes[idx]
        depths = [r.depth_m for r in rows]
        vals = [getattr(r, ind) for r in rows if getattr(r, ind) is not None]
        matched_depths = [d for d, v in zip(depths, [getattr(r, ind) for r in rows]) if v is not None]
        if matched_depths and len(matched_depths) == len(vals):
            ax.scatter(matched_depths, vals, s=4, alpha=0.4, color=COLORS[idx])
            # 异常点高亮
            anom = [a for a in anomalies if a.get("indicator") == sn]
            if anom:
                ad = [a["depth"] for a in anom]
                av = [a["value"] for a in anom]
                ax.scatter(ad, av, s=20, c='red', marker='x', alpha=0.9, label='异常')
        ax.set_xlabel('深度 (m)', fontsize=8)
        title = INDICATOR_CN.get(ind, ind)
        ax.set_title(title, fontsize=9)
        ax.invert_xaxis()
        ax.tick_params(labelsize=7)
    fig.suptitle(f'{task_name} - 深度剖面图', fontsize=12, fontweight='bold', y=1.01)
    plt.tight_layout()
    return _fig_to_bytes(fig)


def make_anomaly_chart(anomalies, stats, task_name):
    """异常分布图：左侧是按指标的柱状图，右侧是按方法的饼图"""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4))

    # 柱状图：按指标
    shorts = ["chl", "odo", "temp", "ph", "turb"]
    labels = [SHORT_CN[s] for s in shorts]
    counts = [stats.get(f"{s}_anomaly_count", 0) for s in shorts]
    bars = ax1.bar(labels, counts, color=COLORS, edgecolor='white')
    for bar, cnt in zip(bars, counts):
        if cnt > 0:
            ax1.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.3, str(cnt),
                     ha='center', fontsize=9)
    ax1.set_title('各指标异常数量', fontsize=11)
    ax1.set_ylabel('异常点数')
    ax1.tick_params(labelsize=8)

    # 饼图：按方法
    tf_count = sum(1 for a in anomalies if a.get("method") == "threshold")
    if_count = sum(1 for a in anomalies if a.get("method") == "isolation_forest")
    if tf_count + if_count > 0:
        pie_data = []
        pie_labels = []
        if tf_count > 0:
            pie_data.append(tf_count)
            pie_labels.append(f'阈值检测 ({tf_count})')
        if if_count > 0:
            pie_data.append(if_count)
            pie_labels.append(f'孤立森林 ({if_count})')
        ax2.pie(pie_data, labels=pie_labels, autopct='%1.1f%%', colors=['#e6a23c', '#409eff'],
                startangle=90, textprops={'fontsize': 9})
        ax2.set_title('检测方法分布', fontsize=11)

    fig.suptitle(f'{task_name} - 异常分析', fontsize=12, fontweight='bold')
    plt.tight_layout()
    return _fig_to_bytes(fig)


# ─── DOCX 生成 ────────────────────────────────────────────

def _set_cell_shading(cell, color):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = tcPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color,
    })
    tcPr.append(shading)


def generate_docx(llm_text: str, chart_images: list[bytes], task_info: dict, stats: dict, task_id: str) -> str:
    os.makedirs(settings.REPORT_DIR, exist_ok=True)

    doc = Document()
    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = '宋体'
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    from docx.oxml import OxmlElement

    # Heading 样式：黑色宋体
    for i in range(1, 4):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.name = '宋体'
        h_style.font.color.rgb = RGBColor(0, 0, 0)
        h_rPr = h_style.element.get_or_add_rPr()
        h_rFonts = h_rPr.find(qn('w:rFonts'))
        if h_rFonts is None:
            h_rFonts = OxmlElement('w:rFonts')
            h_rPr.append(h_rFonts)
        h_rFonts.set(qn('w:eastAsia'), '宋体')

    # ── 封面 ──
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('水质巡航分析报告')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.name = '宋体'
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {qn('w:eastAsia'): '宋体'})
    rPr.append(rFonts)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"—— {task_info.get('reservoir_name', '未知水库')} ——")
    run.font.size = Pt(16)
    run.font.name = '宋体'
    rPr2 = run._r.get_or_add_rPr()
    rFonts2 = rPr2.makeelement(qn('w:rFonts'), {qn('w:eastAsia'): '宋体'})
    rPr2.append(rFonts2)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"采样时间：{task_info.get('created_at', '')}").font.size = Pt(12)
    doc.add_paragraph()
    info2 = doc.add_paragraph()
    info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info2.add_run(f"采样点数：{task_info.get('total_points', 0)}　　深度层：{task_info.get('depth_layers', '')}").font.size = Pt(12)

    doc.add_page_break()

    # ── 目录 ──
    toc_heading = doc.add_heading('目  录', level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_items = [
        '一、指标统计汇总',
        '二、深度剖面图',
        '三、异常分析',
        '四、总体水质评价',
        '五、异常区域分析',
        '六、历史案例对比',
        '七、改善建议',
        '附录：数据说明',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(6)
        p.runs[0].font.size = Pt(12)

    doc.add_page_break()

    # ── 一、指标统计汇总 ──
    doc.add_heading('一、指标统计汇总', level=1)
    shorts = ["chl", "odo", "temp", "ph", "turb"]
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['指标', '均值', '标准差', '最小值', '最大值', '异常率', '异常点数']
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_shading(cell, '333333')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = '宋体'
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.makeelement(qn('w:rFonts'), {qn('w:eastAsia'): '宋体'})
            rPr.append(rFonts)

    for s in shorts:
        row = table.add_row()
        vals = [SHORT_CN[s], f"{stats[f'{s}_mean']:.2f}", f"{stats[f'{s}_std']:.2f}",
                f"{stats[f'{s}_min']:.2f}", f"{stats[f'{s}_max']:.2f}",
                f"{stats[f'{s}_anomaly_rate']:.1%}", str(stats[f'{s}_anomaly_count'])]
        for j, v in enumerate(vals):
            cell = row.cells[j]
            cell.text = v
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
            if j == 5 and stats[f'{s}_anomaly_rate'] > 0.1:
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(0xf5, 0x6c, 0x6c)

    doc.add_paragraph()

    # ── 二、深度剖面图 ──
    doc.add_heading('二、深度剖面图', level=1)
    p = doc.add_paragraph('下图为各指标在不同深度的取值分布，红色叉号标注异常点位。')
    if len(chart_images) >= 1 and chart_images[0]:
        doc.add_picture(chart_images[0], width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ── 三、异常分析 ──
    doc.add_heading('三、异常分析', level=1)
    if len(chart_images) >= 2 and chart_images[1]:
        doc.add_picture(chart_images[1], width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ── LLM 生成的内容 ──
    sections = re.split(r'##SECTION:(.+?)##', llm_text)
    # sections[0] is text before first marker (usually empty)
    section_map = {}
    for i in range(1, len(sections), 2):
        title = sections[i].strip()
        body = sections[i + 1].strip() if i + 1 < len(sections) else ''
        section_map[title] = body

    heading_map = {
        '总体水质评价': ('四、总体水质评价', 1),
        '异常区域分析': ('五、异常区域分析', 1),
        '历史案例对比': ('六、历史案例对比', 1),
        '改善建议': ('七、改善建议', 1),
    }

    for key, (heading, level) in heading_map.items():
        doc.add_heading(heading, level=level)
        content = section_map.get(key, '')
        if content:
            # Split into paragraphs
            paragraphs = [p.strip() for p in content.split('\n') if p.strip()]
            for para_text in paragraphs:
                # Remove any remaining markdown markers
                cleaned = re.sub(r'\*\*|__|#{1,4}\s*|[-*]\s+', '', para_text)
                if cleaned:
                    p = doc.add_paragraph(cleaned)
                    p.paragraph_format.first_line_indent = Cm(0.7)

    # ── 附录 ──
    doc.add_page_break()
    doc.add_heading('附录：数据说明', level=1)
    doc.add_paragraph(f"本报告由水质三维智能监测与分析系统自动生成。")
    doc.add_paragraph(f"数据来源：CSV 文件上传，经 IDW 空间插值后进行统计分析。")
    doc.add_paragraph(f"异常检测方法：统计阈值法（基于标准限值）+ 孤立森林（Isolation Forest）无监督检测。")
    doc.add_paragraph(f"大模型驱动：使用 DeepSeek 大语言模型对统计结果进行智能解读和报告撰写。")
    p = doc.add_paragraph(f"报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x90, 0x93, 0x99)

    path = os.path.join(settings.REPORT_DIR, f"{task_id}.docx")
    doc.save(path)
    return path




def generate(task_info, rows, anomalies, similar=None, progress_callback=None):
    """主流程：统计 → 图表+LLM并行 → DOCX。
    progress_callback(progress: int, phase: str) — 用于外部进度追踪。"""
    task_name = task_info.get('reservoir_name', '') or task_info.get('id', '')[:8]

    if progress_callback:
        progress_callback(5, "正在计算统计指标...")
    stats = compute_stats(rows, anomalies)

    if progress_callback:
        progress_callback(10, "正在构建分析提示...")
    prompt = build_prompt(task_info, stats, anomalies, similar or [])

    # 图表 + LLM 并行（三者无依赖，LLM 是瓶颈）
    if progress_callback:
        progress_callback(15, "正在生成图表与AI分析（预计30-60秒）...")
    from concurrent.futures import ThreadPoolExecutor, as_completed, TimeoutError as FutTimeout
    chart_images = [None, None]
    llm_text = ""
    chart_done = [False, False]
    with ThreadPoolExecutor(max_workers=3) as pool:
        futs = {
            pool.submit(make_depth_profile_chart, rows, anomalies, task_name): "chart1",
            pool.submit(make_anomaly_chart, anomalies, stats, task_name): "chart2",
            pool.submit(call_llm, prompt): "llm",
        }
        for fut in as_completed(futs):
            key = futs[fut]
            try:
                timeout_s = 480 if key == "llm" else 60
                if key == "chart1":
                    chart_images[0] = fut.result(timeout=timeout_s)
                    chart_done[0] = True
                    if progress_callback and chart_done[1]:
                        progress_callback(25, "图表完成，等待AI分析...")
                elif key == "chart2":
                    chart_images[1] = fut.result(timeout=timeout_s)
                    chart_done[1] = True
                    if progress_callback and chart_done[0]:
                        progress_callback(25, "图表完成，等待AI分析...")
                elif key == "llm":
                    llm_text = fut.result(timeout=timeout_s)
                    if progress_callback:
                        progress_callback(80, "AI分析完成，正在生成文档...")
            except FutTimeout:
                logger.warning(f"并行任务 {key} 超时，跳过")
            except Exception as e:
                logger.warning(f"并行任务 {key} 失败: {e}")

    if progress_callback:
        progress_callback(85, "正在生成Word文档...")
    docx_path = generate_docx(llm_text, chart_images, task_info, stats, task_info["id"])

    if progress_callback:
        progress_callback(95, "正在保存文件...")

    return {
        "docx_path": docx_path,
        "llm_text": llm_text,
        "stats": stats,
        "chart1_bytes": chart_images[0].getvalue() if chart_images[0] else None,
        "chart2_bytes": chart_images[1].getvalue() if chart_images[1] else None,
        "chart1_present": chart_images[0] is not None,
        "chart2_present": chart_images[1] is not None,
    }
